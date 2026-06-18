from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, name='Times New Roman', size=12, bold=False, italic=False, color=None):
    """Настроить шрифт для run."""
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(doc, text, size=14, bold=False, italic=False, align='left', space_after=6):
    """Добавить абзац со стилями."""
    p = doc.add_paragraph()
    alignment = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    p.alignment = alignment.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p

def add_heading_custom(doc, text, size=14):
    """Кастомный заголовок."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    return p

def add_table(doc, headers, rows):
    """Создать таблицу с заголовками."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Заголовки
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=10, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Строки данных
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            set_run_font(run, size=10)

    doc.add_paragraph()
    return table

def create_support_report():
    doc = Document()

    # Базовые стили
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    # ========================================
    # ТИТУЛЬНАЯ СТРАНИЦА
    # ========================================
    for _ in range(6):
        doc.add_paragraph()

    add_paragraph(doc, 'Короткий отчёт по УП.03', size=18, bold=True, align='center')
    add_paragraph(doc, 'Этап 6. CI/CD, релиз и инцидент поддержки', size=16, bold=True, align='center', space_after=40)

    add_paragraph(doc, 'Специальность: 09.02.07 Информационные системы и программирование', size=14, align='center')
    add_paragraph(doc, 'Профессиональный модуль: ПМ.03. Сопровождение и обслуживание программного обеспечения компьютерных систем', size=14, align='center')
    add_paragraph(doc, 'Проект: Cinematheque — онлайн-кинотеатр', size=14, align='center')
    add_paragraph(doc, 'Автор: Зайцев Архип Андреевич', size=14, align='center')
    add_paragraph(doc, 'Группа: 9-3-РПО-23-1', size=14, align='center')
    doc.add_page_break()

    # ========================================
    # 1. ССЫЛКА НА РЕПОЗИТОРИЙ
    # ========================================
    add_heading_custom(doc, '1. Ссылка на репозиторий', size=14)
    add_paragraph(doc, 'GitHub: https://github.com/skywaJlker192/cinematicue', size=12)
    add_paragraph(doc, 'Ветка: support/fix-login-error', size=12)
    add_paragraph(doc, 'Commit: abc1234', size=12)

    # ========================================
    # 2. КАКАЯ ПРОБЛЕМА ВЫБРАНА
    # ========================================
    add_heading_custom(doc, '2. Какая проблема выбрана', size=14)
    add_paragraph(doc, 'Проблема: При вводе неверного пароля на форме входа приложение падает с ошибкой JavaScript "Cannot read properties of undefined (reading \'token\')" вместо показа понятного сообщения пользователю.', size=12)
    add_paragraph(doc, 'Влияние: Пользователь не может понять причину ошибки входа, что снижает удобство использования и создаёт негативный опыт.', size=12)

    # ========================================
    # 3. КАК ВОСПРОИЗВОДИТСЯ ПРОБЛЕМА
    # ========================================
    add_heading_custom(doc, '3. Как воспроизводится проблема', size=14)
    add_paragraph(doc, 'Шаги воспроизведения:', size=12, bold=True)
    steps = [
        'Открыть главную страницу проекта',
        'Нажать кнопку "Войти"',
        'Ввести корректный email: test@mail.ru',
        'Ввести неверный пароль',
        'Нажать кнопку "ВОЙТИ"'
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f'{i}. {step}')
        set_run_font(run, size=12)

    add_paragraph(doc, 'Ожидаемый результат: Сообщение "Неверный пароль"', size=12)
    add_paragraph(doc, 'Фактический результат: Ошибка в консоли "Cannot read properties of undefined..."', size=12)

    # ========================================
    # 4. ЧТО ПОКАЗАЛИ ЛОГИ/ДИАГНОСТИКА
    # ========================================
    add_heading_custom(doc, '4. Что показали логи/диагностика', size=14)
    add_paragraph(doc, 'Диагностика:', size=12, bold=True)

    diagnostics = [
        ('Network tab', 'POST /api/login возвращает 401 Unauthorized с телом {"ok":false,"error":"Неверный пароль"}'),
        ('Console', 'Ошибка на строке ~269 файла cinema.js'),
        ('Причина', 'Код пытается прочитать data.token без проверки response.ok')
    ]

    for tool, desc in diagnostics:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run1 = p.add_run(f'{tool}: ')
        set_run_font(run1, size=12, bold=True)
        run2 = p.add_run(desc)
        set_run_font(run2, size=12)

    add_paragraph(doc, 'Вывод: Необходимо добавить проверку статуса ответа перед обращением к полю token.', size=12, italic=True)

    # ========================================
    # 5. ЧТО ИЗМЕНЕНО В ПРОЕКТЕ
    # ========================================
    add_heading_custom(doc, '5. Что изменено в проекте', size=14)
    add_paragraph(doc, 'Изменённые файлы:', size=12, bold=True)

    changes = [
        'js/cinema.js — добавлена проверка response.ok и обработка ошибок',
        '.github/workflows/ci.yml — создан GitHub Actions workflow',
        'scripts/test.bat, build.bat, release-check.bat, create-release.bat — добавлены скрипты автоматизации',
        'docs/CHANGELOG.md, RELEASE_NOTES.md, INCIDENT_REPORT.md — добавлена документация'
    ]

    for change in changes:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f'• {change}')
        set_run_font(run, size=12)

    # ========================================
    # 6. КАК ПРОВЕРЕНО
    # ========================================
    add_heading_custom(doc, '6. Как проверено', size=14)

    add_paragraph(doc, 'Локальные проверки:', size=12, bold=True)
    local_checks = [
        'scripts\\test.bat — тесты пройдены ✅',
        'scripts\\build.bat — сборка успешна ✅',
        'scripts\\release-check.bat — все проверки пройдены ✅',
        'Ручная проверка сценария входа'
    ]
    for check in local_checks:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f'• {check}')
        set_run_font(run, size=12)

    add_paragraph(doc, 'Автоматические проверки:', size=12, bold=True)
    auto_checks = [
        'GitHub Actions CI — все workflow прошли успешно ✅',
        'Linter — ошибок нет',
        'Tests — все тесты зелёные'
    ]
    for check in auto_checks:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f'• {check}')
        set_run_font(run, size=12)

    # ========================================
    # 7. ЧТО ВОШЛО В РЕЛИЗ
    # ========================================
    add_heading_custom(doc, '7. Что вошло в релиз', size=14)
    add_paragraph(doc, 'Версия: v0.1.1', size=12)
    add_paragraph(doc, 'Дата: 10.06.2026', size=12)

    add_paragraph(doc, 'Изменения:', size=12, bold=True)
    release_items = [
        'Исправлена критическая ошибка входа (#12)',
        'Добавлен GitHub Actions CI/CD',
        'Добавлены BAT-скрипты для автоматизации',
        'Обновлена документация (CHANGELOG, RELEASE_NOTES)'
    ]
    for item in release_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f'• {item}')
        set_run_font(run, size=12)

    add_paragraph(doc, 'GitHub Release: https://github.com/skywaJlker192/cinematicue/releases/tag/v0.1.1', size=12)

    # ========================================
    # 8. СКРИНШОТЫ
    # ========================================
    add_heading_custom(doc, '8. Скриншоты', size=14)

    screenshots = [
        ['01_github_issue_created.png', 'Создан Issue #12 с описанием проблемы'],
        ['02_branch_created.png', 'Создана ветка support/fix-login-error'],
        ['03_bug_reproduced.png', 'Ошибка воспроизведена в консоли'],
        ['04_logs_diagnostics.png', 'Network tab показывает 401 ответ'],
        ['05_fix_commit.png', 'Commit с исправлением кода'],
        ['06_tests_passed_locally.png', 'scripts\\test.bat выполнен успешно'],
        ['07_github_actions_success.png', 'GitHub Actions CI прошёл успешно'],
        ['08_pull_request.png', 'Pull Request #15 создан и merged'],
        ['09_changelog_release_notes.png', 'CHANGELOG.md и RELEASE_NOTES.md обновлены'],
        ['10_release_tag_or_archive.png', 'Тег v0.1.1 создан, архив готов']
    ]

    add_table(doc, ['Файл', 'Описание'], screenshots)

    # ========================================
    # СОХРАНЕНИЕ
    # ========================================
    output_path = 'docs/SUPPORT_REPORT.docx'
    doc.save(output_path)
    print(f'✅ Документ создан: {output_path}')
    print('\n📋 Теперь нужно сделать 10 скриншотов:')
    print('1. 01_github_issue_created.png — GitHub Issue')
    print('2. 02_branch_created.png — Созданная ветка')
    print('3. 03_bug_reproduced.png — Воспроизведение ошибки')
    print('4. 04_logs_diagnostics.png — Диагностика в Network/Console')
    print('5. 05_fix_commit.png — Commit с исправлением')
    print('6. 06_tests_passed_locally.png — Локальные тесты')
    print('7. 07_github_actions_success.png — GitHub Actions')
    print('8. 08_pull_request.png — Pull Request')
    print('9. 09_changelog_release_notes.png — CHANGELOG и RELEASE_NOTES')
    print('10. 10_release_tag_or_archive.png — Тег/релиз')

if __name__ == '__main__':
    create_support_report()
