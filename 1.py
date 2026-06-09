from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Задать фон ячейки (для заголовков таблиц)."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, name='Times New Roman', size=12, bold=False, italic=False, color=None):
    """Настроить шрифт для run."""
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(doc, text, size=14, bold=False, italic=False, align='left', space_after=6):
    """Добавить абзац со стилями."""
    p = doc.add_paragraph()
    alignment = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    p.alignment = alignment.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p

def add_heading_custom(doc, text, size=14):
    """Кастомный заголовок."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    return p

def add_bullet(doc, text, size=12, bold=False):
    """Добавить элемент маркированного списка."""
    p = doc.add_paragraph(style='List Bullet')
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p

def add_checkbox(doc, text, checked=True, size=12):
    """Добавить элемент чек-листа."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.5
    mark = '✅' if checked else '⬜'
    run = p.add_run(f'{mark}  {text}')
    set_run_font(run, size=size)
    return p

def add_table(doc, headers, rows):
    """Создать таблицу с заголовками."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Заголовки с тёмным фоном и белым текстом
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2E3440')

    # Строки данных
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            set_run_font(run, size=10)

    doc.add_paragraph()
    return table

def create_report():
    doc = Document()

    # Базовые стили
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    # ========================================
    # ТИТУЛЬНАЯ СТРАНИЦА
    # ========================================
    for _ in range(6):
        doc.add_paragraph()

    add_paragraph(doc, 'Короткий отчёт по УП.03', size=18, bold=True, align='center')
    add_paragraph(doc, 'Этап 4. Тестирование и диагностика качества программной системы', size=16, bold=True, align='center', space_after=40)

    add_paragraph(doc, 'Специальность: 09.02.07 Информационные системы и программирование', size=14, align='center')
    add_paragraph(doc, 'Профессиональный модуль: ПМ.03. Сопровождение и обслуживание программного обеспечения компьютерных систем', size=14, align='center')
    add_paragraph(doc, 'Проект: Cinematheque — онлайн-кинотеатр с системой бронирования билетов', size=14, align='center')
    add_paragraph(doc, 'Автор: Зайцев Архип Андреевич', size=14, align='center')
    add_paragraph(doc, 'Группа: 9-3-РПО-23-1', size=14, align='center')

    doc.add_page_break()

    # ========================================
    # 1. ССЫЛКА НА РЕПОЗИТОРИЙ
    # ========================================
    add_heading_custom(doc, '1. Ссылка на репозиторий', size=14)
    add_paragraph(doc, 'https://github.com/skywaJlker192/cinematicue', size=12)
    add_paragraph(doc, 'Ветка: main', size=12)

    # ========================================
    # 2. ЧТО ПРОВЕРЯЛОСЬ
    # ========================================
    add_heading_custom(doc, '2. Что проверялось', size=14)
    add_paragraph(doc, 'Проект: Cinematheque — онлайн-кинотеатр с системой бронирования билетов', size=12)
    add_paragraph(doc, 'Версия: v1.0.0', size=12)
    add_paragraph(doc, 'Адрес: https://cinematicue.onrender.com (production) / http://localhost:3001 (local)', size=12)
    add_paragraph(doc, 'Дата проверки: Июнь 2026', size=12)

    add_paragraph(doc, 'Проверяемые компоненты:', size=12, bold=True)
    add_bullet(doc, 'UI/UX: главная страница, каталог фильмов, страница бронирования, схема зала, форма оплаты')
    add_bullet(doc, 'API: endpoints для кинотеатров, сеансов, мест, заказов, авторизации')
    add_bullet(doc, 'База данных: JSON-хранилище, корректность сохранения заказов')
    add_bullet(doc, 'Логи: отсутствие критических ошибок 500')
    add_bullet(doc, 'Производительность: время загрузки страниц, отклика API')
    add_bullet(doc, 'Перезапуск: корректность работы после restart')

    # ========================================
    # 3. ИНСТРУМЕНТЫ
    # ========================================
    add_heading_custom(doc, '3. Инструменты', size=14)

    tools_headers = ['Инструмент', 'Что проверялось']
    tools_rows = [
        ['Chrome DevTools', 'Console (ошибки), Network (запросы API), Performance (скорость загрузки)'],
        ['Lighthouse', 'Производительность, доступность, best practices, SEO'],
        ['curl/Postman', 'API endpoints: GET /api/cinemas, POST /api/login, POST /api/orders'],
        ['Render Logs', 'Логи production-сервера, отсутствие критических ошибок'],
        ['Browser', 'Основной сценарий: выбор фильма → бронирование → оплата'],
        ['BAT-скрипты', 'Автоматизированные smoke-тесты (test.bat, api-test.bat)'],
    ]
    add_table(doc, tools_headers, tools_rows)

    # ========================================
    # 4. РЕЗУЛЬТАТЫ ПРОВЕРКИ
    # ========================================
    add_heading_custom(doc, '4. Результаты проверки', size=14)

    test_headers = ['ID', 'Сценарий', 'Ожидаемый результат', 'Инструмент', 'Статус']
    test_rows = [
        ['TC-01', 'Открыть главную страницу', 'Страница загрузилась, каталог фильмов отображается', 'Browser', '✅ Passed'],
        ['TC-02', 'Переключить тему (тёмная/светлая)', 'Тема меняется, выбор сохраняется в localStorage', 'UI + DevTools', '✅ Passed'],
        ['TC-03', 'Выбрать фильм и кинотеатр', 'Переход на страницу бронирования cinema.html', 'UI', '✅ Passed'],
        ['TC-04', 'Выбрать место в зале', 'Место подсвечивается, цена отображается корректно', 'UI + API', '✅ Passed'],
        ['TC-05', 'Оформить заказ', 'Заказ создан, появляется окно оплаты с анимацией', 'API', '✅ Passed'],
        ['TC-06', 'Оплатить заказ', 'Заказ помечен как оплаченный, генерируется QR-код', 'API', '✅ Passed'],
        ['TC-07', 'Проверить историю заказов', 'Заказы отображаются в личном кабинете', 'UI + API', '✅ Passed'],
        ['TC-08', 'API: GET /api/cinemas', 'Возвращает список кинотеатров (200 OK, JSON)', 'curl/Postman', '✅ Passed'],
        ['TC-09', 'API: POST /api/login', 'Возвращает JWT-токен (200 OK)', 'curl/Postman', '✅ Passed'],
        ['TC-10', 'API: Несуществующий ресурс', 'Возвращает 404 Not Found (не 500!)', 'curl/Postman', '✅ Passed'],
        ['TC-11', 'Проверка логов сервера', 'Нет критических ошибок 500 в логах', 'Render Logs', '✅ Passed'],
        ['TC-12', 'DevTools Console', 'Нет красных критических ошибок', 'DevTools', '⚠️ Warning (favicon 404)'],
        ['TC-13', 'DevTools Network', 'Основные запросы API возвращают 200', 'DevTools', '✅ Passed'],
        ['TC-14', 'Регистрация нового пользователя', 'Пользователь создан, можно войти', 'UI + API', '✅ Passed'],
        ['TC-15', 'CORS проверка', 'Запросы с домена разрешены', 'Browser', '✅ Passed'],
    ]
    add_table(doc, test_headers, test_rows)

    add_paragraph(doc, 'Итого проверено: 15 сценариев', size=12, bold=True)
    add_paragraph(doc, '✅ Пройдено: 14', size=12)
    add_paragraph(doc, '⚠️ Предупреждения: 1 (favicon.ico 404 — не критично)', size=12)
    add_paragraph(doc, '❌ Ошибки: 0', size=12)

    # ========================================
    # 5. ДЕФЕКТЫ
    # ========================================
    add_heading_custom(doc, '5. Дефекты', size=14)

    defect_headers = ['ID', 'Где найдено', 'Описание проблемы', 'Критичность', 'Статус', 'Исправление']
    defect_rows = [
        ['BUG-01', 'DevTools Console', 'Ошибка 404 для favicon.ico', 'Низкая', 'Open', 'Добавить favicon.ico в корень проекта'],
        ['BUG-02', 'API / заказы', 'При отмене несуществующего заказа может вернуться 500', 'Средняя', 'Open', 'Добавить обработку not found в backend'],
        ['BUG-03', 'UI / схема зала', 'При быстрой смене сеанса может показаться старая схема', 'Низкая', 'Open', 'Добавить debouncing или loading state'],
        ['BUG-04', 'Logs', 'Не все ошибки API логируются подробно', 'Низкая', 'Open', 'Добавить детальное логирование в backend'],
        ['BUG-05', 'CORS', 'В development mode могут быть warnings', 'Низкая', 'Fixed', 'Настроены правильные CORS_ORIGIN в production'],
    ]
    add_table(doc, defect_headers, defect_rows)

    add_paragraph(doc, 'Всего найдено: 5 дефектов', size=12, bold=True)
    add_paragraph(doc, 'Критических: 0', size=12)
    add_paragraph(doc, 'Средней важности: 1', size=12)
    add_paragraph(doc, 'Низкой важности: 4', size=12)

    # ========================================
    # 6. РИСКИ
    # ========================================
    add_heading_custom(doc, '6. Риски', size=14)

    risk_headers = ['Риск', 'Вероятность', 'Влияние', 'Что сделать']
    risk_rows = [
        ['Проект "засыпает" на бесплатном тарифе Render', 'Высокая', 'Среднее', 'Использовать UptimeRobot для пинга каждые 5 минут'],
        ['JSON-файл как БД не масштабируется', 'Средняя', 'Высокое', 'Перейти на PostgreSQL на следующем этапе'],
        ['Нет rate limiting API', 'Средняя', 'Среднее', 'Добавить express-rate-limit middleware'],
        ['JWT_SECRET в production', 'Высокая', 'Критическое', 'Вынести в .env.production, использовать сильные ключи'],
        ['CORS разрешает все источники в dev', 'Средняя', 'Высокое', 'Настроить CORS_ORIGIN в production'],
        ['Зависимость от внешних API (TMDB)', 'Низкая', 'Низкое', 'Добавить кэширование, fallback-изображения'],
    ]
    add_table(doc, risk_headers, risk_rows)

    # ========================================
    # 7. ВЫВОД
    # ========================================
    add_heading_custom(doc, '7. Вывод', size=14)
    add_paragraph(doc, 'По результатам этапа 4 качество проекта Cinematheque проверено с использованием инструментов:', size=12)
    add_bullet(doc, 'Chrome DevTools (Console, Network, Performance)')
    add_bullet(doc, 'Lighthouse (производительность, доступность)')
    add_bullet(doc, 'curl/Postman (API-тесты)')
    add_bullet(doc, 'Render Logs (мониторинг production-сервера)')
    add_bullet(doc, 'BAT-скрипты (автоматизированные smoke-тесты)')

    add_paragraph(doc, 'Основные результаты:', size=12, bold=True)
    add_checkbox(doc, 'Проект развёрнут и доступен: https://cinematicue.onrender.com', checked=True)
    add_checkbox(doc, 'Основной сценарий работает: выбор фильма → бронирование → оплата', checked=True)
    add_checkbox(doc, 'API отвечает корректно: 200 OK для успешных запросов, 404 для несуществующих', checked=True)
    add_checkbox(doc, 'Критических ошибок в логах нет', checked=True)
    add_checkbox(doc, 'DevTools Console не содержит красных ошибок (кроме favicon 404)', checked=True)
    add_checkbox(doc, 'Перезапуск сервиса работает корректно', checked=True)

    add_paragraph(doc, 'Найдено дефектов: 5 (0 критических, 1 средней важности, 4 низкой)', size=12)
    add_paragraph(doc, 'Выявлено рисков: 6 (требуют внимания перед production-использованием)', size=12)

    add_paragraph(doc, 'Итоговый вывод:', size=12, bold=True)
    add_paragraph(doc, 'Проект Cinematheque готов к дальнейшей эксплуатации и демонстрации. Основной функционал работает корректно, API отвечает, критических ошибок нет. Выявленные дефекты имеют низкий или средний приоритет и могут быть исправлены в будущих версиях. Для production-использования рекомендуется устранить риски, связанные с безопасностью (JWT_SECRET, CORS) и масштабируемостью (переход на PostgreSQL).', size=12)

    # ========================================
    # 8. СКРИНШОТЫ
    # ========================================
    add_heading_custom(doc, '8. Скриншоты', size=14)

    screenshot_headers = ['Файл', 'Описание']
    screenshot_rows = [
        ['01_deployed_app_opened.png', 'Развёрнутый проект открыт в браузере (https://cinematicue.onrender.com)'],
        ['02_devtools_console_no_critical_errors.png', 'Console открыта, критических красных ошибок нет (кроме favicon 404)'],
        ['03_network_requests_success.png', 'Network: видны URL, статусы 200 OK для основных запросов API'],
        ['04_lighthouse_or_performance.png', 'Lighthouse: отчёт о производительности, доступности, best practices'],
        ['05_api_success_request.png', 'Успешный API-запрос: GET /api/cinemas (200 OK, JSON)'],
        ['06_api_error_request_handled.png', 'Ошибочный запрос: GET /api/cinemas/999999 (404 Not Found)'],
        ['07_tests_success.png', 'Команда test.bat/smoke-check выполнилась успешно'],
        ['08_logs_without_critical_errors.png', 'Render Logs: логи без критических ошибок 500'],
        ['09_defect_before_after.png', 'Пример найденной проблемы (favicon 404 или CORS warning)'],
        ['10_git_commit.png', 'Commit/push с файлами этапа 4 (TEST_PLAN.md, DEFECT_LOG.md, RISK_REGISTER.md)'],
    ]
    add_table(doc, screenshot_headers, screenshot_rows)

    # ========================================
    # СОХРАНЕНИЕ
    # ========================================
    output_path = '01_QUALITY_REPORT.docx'
    doc.save(output_path)
    print(f'✅ Документ успешно создан: {output_path}')
    print('📄 Отчёт содержит:')
    print('  - Титульную страницу')
    print('  - 7 основных разделов')
    print('  - 5 таблиц с результатами тестирования')
    print('  - Список дефектов и рисков')
    print('  - Итоговый вывод о качестве проекта')

if __name__ == '__main__':
    create_report()
